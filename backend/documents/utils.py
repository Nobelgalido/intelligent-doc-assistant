import PyPDF2
from docx import Document as DocxDocument
from typing import Tuple, List
import re

class DocumentProcessor:
    """Utility class for document text extraction and processing."""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> tuple[str, int]:
        """Extract text and page count from a PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
                    
            return text.strip(), page_count
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
        
    @staticmethod
    def extract_text_from_docx(file_path: str) -> tuple[str, int]:
        """Extract text and page count from a DOCX file."""
        text = ""
        try:
            doc = DocxDocument(file_path)
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text])
            page_count = len(doc.paragraphs) // 20  # Rough estimate: 20 paragraphs per page
            return text, max(1,page_count)
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Chunk text into smaller pieces with specified size and overlap."""
        text = re.sub(r'\s+', ' ', text).strip()
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size
            
            if end >= len(text):
                last_period = text.rfind('.', start, end)
                if last_period > start:
                    end = last_period + 1
                    
            chunk = text[start:end].strip()
            
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            
        return chunks
    
    @staticmethod
    def count_words(text: str) -> int:
        """Count the number of words in the text."""
        return len(text.split())